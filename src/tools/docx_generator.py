"""DOCX generation tool for roadmap Phase 13."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from src.core.storage import backup_file
from src.schemas.citation import ReferenceList
from src.schemas.project import Project, ProjectArtifact
from src.tools.base import BaseTool, ToolRequest, ToolResponse

__all__ = ["DocxGenerationRequest", "DocxGenerationResponse", "DocxGenerationTool"]


class DocxGenerationRequest(ToolRequest):
    project: Project
    draft: str
    reference_list: ReferenceList | None = None
    citation_audit_passed: bool
    fact_audit_passed: bool


class DocxGenerationResponse(ToolResponse):
    docx_path: str | None = None
    backup_path: str | None = None


class DocxGenerationTool(BaseTool[DocxGenerationRequest, DocxGenerationResponse]):
    response_model = DocxGenerationResponse
    tool_name = "docx_generation"

    def _execute(self, request: DocxGenerationRequest) -> DocxGenerationResponse:
        if not (request.citation_audit_passed and request.fact_audit_passed):
            return DocxGenerationResponse.failure(
                error_code="AUDIT_NOT_PASSED",
                error_message="DOCX generation requires passed citation and fact audits",
            )

        path = request.project.artifact_path(ProjectArtifact.FINAL_DOCX)
        backup = backup_file(path, root=request.project.directory)
        doc = Document()
        self._add_markdown(doc, request.draft)
        if request.reference_list and request.reference_list.entries:
            doc.add_heading("References", level=1)
            for entry in request.reference_list.entries:
                doc.add_paragraph(entry.formatted)
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        return DocxGenerationResponse(docx_path=str(path), backup_path=str(backup) if backup else None)

    @staticmethod
    def _add_markdown(doc: Document, text: str) -> None:
        for raw in text.splitlines():
            line = raw.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=2)
            elif line.startswith("> "):
                paragraph = doc.add_paragraph(line[2:].strip())
                paragraph.style = "Quote"
            else:
                doc.add_paragraph(line)
