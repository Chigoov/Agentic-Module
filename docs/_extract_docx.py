import docx

d = docx.Document("Master_Task_Setiap_Fase_AUTONOMI_AGENTIC_ILMIAH.docx")
out = []

out.append("=== PARAGRAPHS ===")
for p in d.paragraphs:
    if p.text.strip():
        out.append(p.text)

out.append("")
out.append("=== TABLES ===")
for ti, t in enumerate(d.tables):
    out.append(f"--- TABLE {ti} ---")
    for r in t.rows:
        cells = [c.text.replace("\n", " / ").strip() for c in r.cells]
        out.append(" | ".join(cells))

text = "\n".join(out)
with open("_master_task_extract.txt", "w", encoding="utf-8") as f:
    f.write(text)
print(f"Wrote {len(text)} chars")
